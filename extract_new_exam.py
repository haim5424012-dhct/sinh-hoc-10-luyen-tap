from pathlib import Path
from docx import Document
import json

path = Path('/home/ubuntu/upload/pasted_file_G1QwMH_Ma_de_101.docx')
doc = Document(path)
paragraphs = []
for i, p in enumerate(doc.paragraphs, 1):
    text = p.text.strip()
    if text:
        paragraphs.append({'index': i, 'style': p.style.name if p.style else '', 'text': text})
tables = []
for ti, table in enumerate(doc.tables, 1):
    rows = []
    for ri, row in enumerate(table.rows, 1):
        rows.append({'row': ri, 'cells': [cell.text.strip() for cell in row.cells]})
    tables.append({'index': ti, 'rows': rows})
out = {'file': path.name, 'paragraphs': paragraphs, 'tables': tables}
Path('/home/ubuntu/sinh-hoc-10-luyen-tap/new_exam_extracted.json').write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding='utf-8')
print({'file': path.name, 'paragraphs': len(paragraphs), 'tables': len(tables)})
