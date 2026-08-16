from pathlib import Path
from docx import Document
import json

inputs = [
    Path('/home/ubuntu/upload/pasted_file_dY3nx7_NHCHKTGKI_SH10_25-26.docx'),
    Path('/home/ubuntu/upload/pasted_file_JpRBcu_Ma_de_101.docx'),
]
output = []
for path in inputs:
    doc = Document(path)
    paragraphs = []
    for i, p in enumerate(doc.paragraphs, 1):
        text = p.text.strip()
        if text:
            paragraphs.append({'index': i, 'style': p.style.name if p.style else '', 'text': text})
    tables = []
    for ti, table in enumerate(doc.tables, 1):
        rows = []
        for ri, row in enumerate(table.rows, 1):
            rows.append({'row': ri, 'cells': [cell.text.strip() for cell in row.cells]})
        tables.append({'index': ti, 'rows': rows})
    output.append({'file': path.name, 'paragraphs': paragraphs, 'tables': tables})
Path('/home/ubuntu/sinh-hoc-10-luyen-tap/extracted_docx.json').write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding='utf-8')
print(json.dumps([{'file': item['file'], 'paragraphs': len(item['paragraphs']), 'tables': len(item['tables'])} for item in output], ensure_ascii=False))
